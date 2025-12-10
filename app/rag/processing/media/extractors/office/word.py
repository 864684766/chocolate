"""
Word内容提取器

使用python-docx提取Word文档的内容。
支持.doc转.docx（使用LibreOffice）。
"""

import logging
import io
import platform
from pathlib import Path
from typing import Dict, Any, List
from .base import OfficeDocumentExtractorBase

logger = logging.getLogger(__name__)

# 配置常量
TABLE_EXTRACTION_ENABLED = True


class WordExtractor(OfficeDocumentExtractorBase):
    """Word内容提取器
    
    使用python-docx提取Word文档的文本、段落、表格等内容。
    对于.doc旧格式，先转换为.docx再处理。
    """
    
    def extract(self, content: bytes, meta: Dict[str, Any]) -> Dict[str, Any]:
        """
        从Word文件中提取内容
        
        用处：从Word文档中提取文本、段落、表格等信息，
        为后续的RAG处理提供结构化数据。
        
        Args:
            content: Word文件的二进制内容
            meta: Word文件的元数据信息
            
        Returns:
            Dict[str, Any]: 提取结果字典，包含paragraphs、tables等
        """
        if not self.is_available():
            logger.error("python-docx不可用，无法提取Word内容")
            raise RuntimeError("python-docx未安装，无法提取Word内容")
        
        # 规范化元数据
        self._normalize_metadata(meta)
        
        try:
            # 检查是否为.doc格式
            filename = meta.get("filename", "")
            if filename.lower().endswith(".doc"):
                logger.warning("检测到.doc格式，尝试转换为.docx")
                try:
                    content = self._convert_doc_to_docx(content)
                except Exception as e:
                    logger.error(f".doc转.docx失败: {str(e)}")
                    raise RuntimeError(f"无法处理.doc格式: {str(e)}") from e
            
            # 打开Word文档
            doc = WordExtractor._open_document(content)
            
            # 提取段落和表格
            paragraphs_data = self._extract_paragraphs(doc)
            tables_data = self._extract_tables(doc) if TABLE_EXTRACTION_ENABLED else []
            
            result = {
                "paragraphs": paragraphs_data,
                "tables": tables_data,
                "total_paragraphs": len(paragraphs_data),
                "total_tables": len(tables_data)
            }
            
            if not self._validate_content(result):
                logger.warning("提取的Word内容验证失败")
            
            return result
            
        except Exception as e:
            logger.error(f"Word提取失败: {str(e)}", exc_info=True)
            raise RuntimeError(f"Word提取失败: {str(e)}") from e
    
    @staticmethod
    def _open_document(content: bytes) -> Any:
        """
        打开Word文档
        
        用处：从二进制内容创建python-docx文档对象。
        
        Args:
            content: Word文件的二进制内容
            
        Returns:
            Document: python-docx文档对象
        """
        from docx import Document
        doc_stream = io.BytesIO(content)
        return Document(doc_stream)
    
    @staticmethod
    def _extract_paragraphs(doc: Any) -> List[Dict[str, Any]]:
        """
        提取所有段落内容
        
        用处：遍历Word文档所有段落，提取每段的文本内容和元数据。
        
        Args:
            doc: python-docx文档对象
            
        Returns:
            List[Dict[str, Any]]: 段落数据列表，每个元素包含text和meta
        """
        paragraphs_data = []
        
        for para_idx, para in enumerate(doc.paragraphs):
            para_text = para.text.strip()
            if not para_text:
                continue
            
            para_meta = WordExtractor._build_paragraph_meta(para_idx, para)
            paragraphs_data.append({
                "text": para_text,
                "meta": para_meta
            })
        
        return paragraphs_data
    
    @staticmethod
    def _build_paragraph_meta(para_idx: int, para: Any) -> Dict[str, Any]:
        """
        构建段落元数据
        
        用处：提取段落的基础信息，这些信息会在分块策略中使用。
        注意：只提取metadata_whitelist中的字段，避免提取无用的字段。
        
        Args:
            para_idx: 段落索引（从0开始）
            para: python-docx段落对象
            
        Returns:
            Dict[str, Any]: 段落元数据字典
        """
        meta = {
            "paragraph_index": para_idx + 1  # 段落索引从1开始
        }
        
        # 检查是否为标题样式（Heading 1-9）
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading"):
            try:
                # 提取标题级别（Heading 1 -> level 1）
                level = int(style_name.split()[-1])
                meta["heading_level"] = level
            except (ValueError, IndexError):
                pass
        
        return meta
    
    @staticmethod
    def _extract_tables(doc: Any) -> List[Dict[str, Any]]:
        """
        提取所有表格内容
        
        用处：遍历Word文档所有表格，提取表格数据。
        
        Args:
            doc: python-docx文档对象
            
        Returns:
            List[Dict[str, Any]]: 表格数据列表，每个元素包含data和meta
        """
        tables_data = []
        
        for table_idx, table in enumerate(doc.tables):
            table_rows = WordExtractor._extract_table_rows(table)
            if not table_rows:
                continue
            
            table_meta = {
                "table_index": table_idx + 1  # 表格索引从1开始
            }
            
            tables_data.append({
                "data": table_rows,
                "meta": table_meta
            })
        
        return tables_data
    
    @staticmethod
    def _extract_table_rows(table: Any) -> List[List[str]]:
        """
        提取表格行数据
        
        用处：从表格对象中提取所有行的数据。
        
        Args:
            table: python-docx表格对象
            
        Returns:
            List[List[str]]: 表格行数据，二维列表
        """
        rows = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            # 过滤空行
            if any(cell for cell in row_data):
                rows.append(row_data)
        return rows
    
    def _get_document_type(self) -> str:
        """
        获取文档类型
        
        Returns:
            str: "word"
        """
        return "word"
    
    def is_available(self) -> bool:
        """
        检查python-docx是否可用
        
        Returns:
            bool: True表示python-docx已安装，False表示未安装
        """
        try:
            import docx  # python-docx
            return True
        except ImportError:
            logger.warning("python-docx未安装，Word提取器不可用")
            return False
    
    @staticmethod
    def _get_libreoffice_command() -> str:
        """
        获取LibreOffice命令名称
        
        用处：根据操作系统返回正确的LibreOffice命令名称。
        Windows上通常是soffice.exe，Linux/Mac上通常是libreoffice。
        
        Returns:
            str: LibreOffice命令名称
        """
        system = platform.system()
        if system == "Windows":
            return "soffice.exe"
        else:
            return "libreoffice"
    
    @staticmethod
    def _get_temp_dir() -> Path:
        """
        获取Word临时文件目录路径
        
        用处：从配置文件中读取临时文件目录路径，在temp_dir下创建word子文件夹。
        
        Returns:
            Path: Word临时文件目录路径
            
        Raises:
            ValueError: 如果配置中的 temp_dir 为空或未配置
            RuntimeError: 如果无法读取配置
        """
        from app.config import get_config_manager
        
        config_manager = get_config_manager()
        media_config = config_manager.get_media_processing_config()
        temp_dir_path = media_config.get("temp_dir", "").strip()
        
        if not temp_dir_path:
            raise ValueError(
                "临时文件目录未配置。请在 config/app_config.json 的 "
                "media_processing.temp_dir 中指定临时文件目录路径。"
            )
        
        base_temp_dir = Path(temp_dir_path).expanduser().resolve()
        if not base_temp_dir.exists():
            base_temp_dir.mkdir(parents=True, exist_ok=True)
            logger.debug(f"Created base temp directory from config: {base_temp_dir}")
        
        # 创建word子文件夹
        word_temp_dir = base_temp_dir / "word"
        if not word_temp_dir.exists():
            word_temp_dir.mkdir(parents=True, exist_ok=True)
            logger.debug(f"Created word temp directory: {word_temp_dir}")
        
        return word_temp_dir
    
    @staticmethod
    def _create_temp_doc_file(content: bytes) -> str:
        """
        创建临时.doc文件
        
        用处：将.doc文件的二进制内容写入临时文件，供LibreOffice转换使用。
        
        Args:
            content: .doc文件的二进制内容
            
        Returns:
            str: 临时文件的绝对路径
        """
        import tempfile
        import os
        
        temp_dir = WordExtractor._get_temp_dir()
        temp_doc_file = tempfile.NamedTemporaryFile(
            delete=False, 
            suffix=".doc",
            dir=str(temp_dir)
        )
        temp_doc_file.write(content)
        temp_doc_file.flush()
        os.fsync(temp_doc_file.fileno())
        temp_doc_path = os.path.abspath(temp_doc_file.name)
        temp_doc_file.close()
        return temp_doc_path
    
    @staticmethod
    def _calculate_docx_path(doc_path: str) -> str:
        """
        计算预期的.docx输出路径
        
        用处：LibreOffice转换时会替换扩展名，而不是追加。
        例如：tmp2y83lik7.doc -> tmp2y83lik7.docx（不是tmp2y83lik7.doc.docx）
        
        Args:
            doc_path: .doc文件的路径
            
        Returns:
            str: 预期的.docx文件路径
        """
        import os
        temp_docx_path = os.path.splitext(doc_path)[0] + ".docx"
        logger.debug(f"临时文件路径: {doc_path}, 预期转换后路径: {temp_docx_path}")
        return temp_docx_path
    
    @staticmethod
    def _run_libreoffice_conversion(doc_path: str) -> None:
        """
        执行LibreOffice转换命令
        
        用处：调用LibreOffice命令行工具将.doc文件转换为.docx格式。
        
        Args:
            doc_path: .doc文件的路径
            
        Raises:
            RuntimeError: 如果转换失败
        """
        import os
        import subprocess
        
        libreoffice_cmd = WordExtractor._get_libreoffice_command()
        logger.debug(f"尝试使用LibreOffice转换: {libreoffice_cmd}, 输入文件: {doc_path}")
        
        result = subprocess.run(
            [
                libreoffice_cmd,
                "--headless",
                "--convert-to", "docx",
                "--outdir", os.path.dirname(doc_path),
                doc_path
            ],
            capture_output=True,
            text=True,
            timeout=30
        )
        
        if result.returncode != 0:
            logger.error(f"LibreOffice转换失败，返回码: {result.returncode}, 错误: {result.stderr}, 输出: {result.stdout}")
            raise RuntimeError(f"LibreOffice转换失败: {result.stderr}")
    
    @staticmethod
    def _build_possible_docx_paths(doc_path: str, expected_docx_path: str) -> list:
        """
        构建可能的.docx文件路径列表
        
        用处：LibreOffice可能使用不同的命名方式，需要尝试多个可能的路径。
        
        Args:
            doc_path: 原始.doc文件路径
            expected_docx_path: 预期的.docx文件路径
            
        Returns:
            list: 可能的.docx文件路径列表
        """
        import os
        
        output_dir = os.path.dirname(doc_path)
        base_name = os.path.splitext(os.path.basename(doc_path))[0]
        possible_paths = [
            expected_docx_path,
            os.path.join(output_dir, base_name + ".docx")
        ]
        
        # 列出输出目录中的所有文件，查找可能的.docx文件
        if os.path.exists(output_dir):
            for file in os.listdir(output_dir):
                if file.endswith(".docx") and file.startswith(base_name):
                    possible_paths.append(os.path.join(output_dir, file))
        
        return possible_paths
    
    @staticmethod
    def _find_converted_file(doc_path: str, expected_docx_path: str) -> str:
        """
        查找转换后的.docx文件
        
        用处：LibreOffice可能使用不同的命名方式，需要查找实际生成的.docx文件。
        
        Args:
            doc_path: 原始.doc文件路径
            expected_docx_path: 预期的.docx文件路径
            
        Returns:
            str: 找到的.docx文件路径
            
        Raises:
            RuntimeError: 如果找不到转换后的文件
        """
        import os
        
        if os.path.exists(expected_docx_path):
            logger.debug(f"LibreOffice转换成功: {doc_path} -> {expected_docx_path}")
            return expected_docx_path
        
        possible_paths = WordExtractor._build_possible_docx_paths(doc_path, expected_docx_path)
        
        for path in possible_paths:
            if os.path.exists(path):
                logger.debug(f"找到转换后的文件: {path}")
                return path
        
        logger.error(f"LibreOffice转换后文件不存在: {expected_docx_path}，已尝试路径: {possible_paths}")
        raise RuntimeError(f"LibreOffice转换后文件不存在: {expected_docx_path}，已尝试路径: {possible_paths}")
    
    @staticmethod
    def _read_converted_file(docx_path: str) -> bytes:
        """
        读取转换后的.docx文件内容
        
        用处：从磁盘读取转换后的.docx文件的二进制内容。
        
        Args:
            docx_path: .docx文件的路径
            
        Returns:
            bytes: .docx文件的二进制内容
            
        Raises:
            RuntimeError: 如果文件不存在
        """
        import os
        
        if not os.path.exists(docx_path):
            logger.error(f"转换后的文件不存在: {docx_path}")
            raise RuntimeError(f"转换后的文件不存在: {docx_path}")
        
        with open(docx_path, "rb") as f:
            return f.read()
    
    @staticmethod
    def _build_libreoffice_error_message(error: Exception) -> str:
        """
        构建LibreOffice未找到的错误消息
        
        用处：当LibreOffice命令未找到时，提供详细的错误提示信息。
        
        Args:
            error: 原始异常对象
            
        Returns:
            str: 格式化的错误消息
        """
        libreoffice_cmd = WordExtractor._get_libreoffice_command()
        return (
            f"未找到LibreOffice，无法转换.doc文件。\n"
            f"请确保已安装LibreOffice，并且命令 '{libreoffice_cmd}' 在系统PATH中。\n"
            f"Windows用户：请将LibreOffice安装目录（如 C:\\Program Files\\LibreOffice\\program）添加到系统PATH。\n"
            f"错误详情: {str(error)}"
        )
    
    @staticmethod
    def _cleanup_temp_files(doc_path: str, docx_path: str) -> None:
        """
        清理临时文件
        
        用处：删除转换过程中创建的临时.doc和.docx文件。
        
        Args:
            doc_path: .doc临时文件路径
            docx_path: .docx临时文件路径
        """
        import os
        
        try:
            if os.path.exists(doc_path):
                os.unlink(doc_path)
            if os.path.exists(docx_path):
                os.unlink(docx_path)
        except Exception as e:
            logger.warning(f"清理临时文件失败: {str(e)}")
    
    @staticmethod
    def _convert_doc_to_docx(content: bytes) -> bytes:
        """
        将.doc格式转换为.docx格式
        
        用处：处理旧格式的Word文档，使用LibreOffice进行转换。
        
        Args:
            content: .doc文件的二进制内容
            
        Returns:
            bytes: .docx文件的二进制内容
            
        Raises:
            RuntimeError: 如果转换失败或LibreOffice不可用
        """
        temp_doc_path = WordExtractor._create_temp_doc_file(content)
        temp_docx_path = WordExtractor._calculate_docx_path(temp_doc_path)
        
        try:
            WordExtractor._run_libreoffice_conversion(temp_doc_path)
            found_docx_path = WordExtractor._find_converted_file(temp_doc_path, temp_docx_path)
            return WordExtractor._read_converted_file(found_docx_path)
        except FileNotFoundError as e:
            error_msg = WordExtractor._build_libreoffice_error_message(e)
            logger.error(error_msg)
            raise RuntimeError(error_msg) from e
        finally:
            WordExtractor._cleanup_temp_files(temp_doc_path, temp_docx_path)
